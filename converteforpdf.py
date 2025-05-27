```
from docx import Document
from datetime import datetime
import os
import re
from docx2pdf import convert 

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            if f'{{{{{key}}}}}' in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if f'{{{{{key}}}}}' in inline[i].text:
                        inline[i].text = inline[i].text.replace(f'{{{{{key}}}}}', val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, val in replacements.items():
                        if f'{{{{{key}}}}}' in paragraph.text:
                            inline = paragraph.runs
                            for i in range(len(inline)):
                                if f'{{{{{key}}}}}' in inline[i].text:
                                    inline[i].text = inline[i].text.replace(f'{{{{{key}}}}}', val)

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]',"", filename)

def main():
    nome_inquilino = input("Nome do inquilino: ")
    
    replacements = {
        'NOME_INQUILINO': nome_inquilino,
        'CPF_INQUILINO': input("CPF do inquilino: "),
        'ENDERECO_INQUILINO': input("Endereço do inquilino: "),
        'CEP_INQUILINO': input("CEP do inquilino: "),
        'TELEFONE_INQUILINO': input("Telefone do inquilino: "),
        'NUMERO_PESSOAS': input("Número máximo de pessoas: "),
        'DATA_ENTRADA': input("Data de entrada (ex: 13/03/2025): "),
        'HORA_ENTRADA': input("Hora de entrada (ex: 10:00h): "),
        'DATA_SAIDA': input("Data de saída (ex: 17/03/2025): "),
        'HORA_SAIDA': input("Hora de saída (ex: 12:00h): "),
        'QUANTIDADE_DIAS': input("Quantidade de diárias: "),
        'VALOR_TOTAL': input("Valor total do aluguel: "),
        'VALOR_SINAL': input("Valor do sinal pago: "),
        'VALOR_PARA_QUITAÇÃO': input("Valor para Quitação: "),
        'DATA_SINAL': input("Data do pagamento do sinal: "),
        'DATA_ATUAL': datetime.today().strftime('%d-%m-%Y'),
    }

    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(current_dir, 'Contrato_Template.docx')

    doc = Document(template_path)
    replace_placeholders(doc, replacements)

    nome_arquivo_limpo = sanitize_filename(nome_inquilino.replace(" ", "_"))
    data_atual = replacements['DATA_ATUAL']

    output_filename = f"Contrato_{nome_arquivo_limpo}_{data_atual}.docx"
    output_path = os.path.join(current_dir, output_filename)

    doc.save(output_path)
    print(f"Contrato preenchido gerado: {output_path}")

    # ✅ Conversão para PDF
    pdf_output_path = output_path.replace('.docx', '.pdf')
    convert(output_path, pdf_output_path)
    print(f"Contrato convertido em PDF: {pdf_output_path}")

if __name__ == '__main__':
    main()

```
